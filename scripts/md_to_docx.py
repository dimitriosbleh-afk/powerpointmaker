"""Convert a unit Markdown document into a Word document.

Written for AUSLAN_1_UNIT_PROMPT.md, which delivers as a Word file. Handles
# / ## headings, \newpage page breaks, markdown tables (with <br> and
**bold** inside cells), - bullets, inline bold, and a Word TOC field.

    python scripts/md_to_docx.py in.md out.docx ["Title"] ["Subtitle"] ["Meta"]

Title defaults to the output filename stem. Subtitle and Meta are optional
title-page lines. There is no pandoc in this environment; python-docx is the
route. QA the result with LibreOffice + PyMuPDF (pdftoppm is not installed):

    soffice --headless --convert-to pdf out.docx
    python -c "import fitz; [p.get_pixmap(dpi=95).save(f'p{i}.png') for i,p in enumerate(fitz.open('out.pdf'))]"
"""
import os
import re
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = sys.argv[1]
OUT = sys.argv[2]
TITLE = sys.argv[3] if len(sys.argv) > 3 else os.path.splitext(os.path.basename(OUT))[0]
SUBTITLE = sys.argv[4] if len(sys.argv) > 4 else ''
META = sys.argv[5] if len(sys.argv) > 5 else ''

USABLE_CM = 18.0  # A4 portrait, 1.5cm margins each side


def add_toc_field(par):
    r = par.add_run()
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'begin')
    r._r.append(fld)
    r2 = par.add_run()
    it = OxmlElement('w:instrText')
    it.set(qn('xml:space'), 'preserve')
    it.text = r'TOC \o "1-2" \h \z \u'
    r2._r.append(it)
    r3 = par.add_run()
    sep = OxmlElement('w:fldChar')
    sep.set(qn('w:fldCharType'), 'separate')
    r3._r.append(sep)
    r4 = par.add_run('Right-click here and choose Update Field to build the contents list.')
    r4.italic = True
    r5 = par.add_run()
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    r5._r.append(end)


def emit_inline(par, text):
    """Write text into a paragraph, honouring **bold** and <br> line breaks."""
    for bi, chunk in enumerate(text.split('<br>')):
        if bi:
            par.add_run().add_break(WD_BREAK.LINE)
        for i, piece in enumerate(re.split(r'\*\*(.+?)\*\*', chunk)):
            if not piece:
                continue
            run = par.add_run(piece)
            if i % 2 == 1:
                run.bold = True


def split_row(line):
    return [c.strip() for c in line.strip().strip('|').split('|')]


def is_sep(line):
    return bool(re.match(r'^\|[\s:|-]+\|$', line.strip())) and '-' in line


CHAR_CM = 0.185   # approx width of one Calibri 11pt character
PAD_CM = 0.42     # left + right cell padding


def col_widths(rows, ncol):
    """Give every column at least the width of its longest unbreakable word,
    then share what is left in proportion to how much text the column holds."""
    floors, weights = [], []
    for c in range(ncol):
        longest_word, longest_cell = 1, 1
        for r in rows:
            txt = r[c].replace('<br>', ' ').replace('**', '')
            longest_cell = max(longest_cell, len(txt))
            for w in txt.split():
                longest_word = max(longest_word, len(w))
        floors.append(min(longest_word * CHAR_CM + PAD_CM, USABLE_CM / 2))
        weights.append(min(longest_cell, 60))

    spare = USABLE_CM - sum(floors)
    if spare <= 0:                       # floors alone overflow: scale them down
        k = USABLE_CM / sum(floors)
        return [Cm(f * k) for f in floors]
    total = sum(weights)
    return [Cm(f + spare * w / total) for f, w in zip(floors, weights)]


def shade(cell, colour):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:color'), 'auto')
    sh.set(qn('w:fill'), colour)
    tcPr.append(sh)


def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement('w:tblHeader')
    el.set(qn('w:val'), 'true')
    trPr.append(el)


doc = Document()

sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
sec.left_margin = sec.right_margin = Cm(1.5)
sec.top_margin = sec.bottom_margin = Cm(1.8)

normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.space_before = Pt(0)

for name, size, colour in (('Heading 1', 20, RGBColor(0x1F, 0x3B, 0x63)),
                           ('Heading 2', 15, RGBColor(0x1F, 0x3B, 0x63))):
    st = doc.styles[name]
    st.font.name = 'Calibri'
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = colour
    st.paragraph_format.space_before = Pt(12)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.keep_with_next = True

# ---- title page ----
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run(TITLE)
r.bold = True
r.font.size = Pt(28)
for text, size in ((SUBTITLE, 16), (META, 12)):
    if not text:
        continue
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text).font.size = Pt(size)
t4 = doc.add_paragraph()
t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t4.add_run('(c) 2026 James Hooke. Confidential. Internal use only. Not for redistribution.')
r.font.size = Pt(9)
r.italic = True
doc.add_paragraph()
h = doc.add_paragraph()
r = h.add_run('Contents')
r.bold = True
r.font.size = Pt(16)
add_toc_field(doc.add_paragraph())
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

lines = open(SRC, encoding='utf-8').read().split('\n')
i = 0
pending_break = False

while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    if stripped == '\\newpage':
        pending_break = True
        i += 1
        continue

    if not stripped:
        i += 1
        continue

    if pending_break:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        pending_break = False

    m = re.match(r'^(#{1,2}) (.+)$', stripped)
    if m:
        doc.add_heading(m.group(2).strip(), level=len(m.group(1)))
        i += 1
        continue

    if stripped.startswith('|'):
        block = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            if not is_sep(lines[i]):
                block.append(split_row(lines[i]))
            i += 1
        ncol = max(len(r) for r in block)
        block = [r + [''] * (ncol - len(r)) for r in block]
        widths = col_widths(block, ncol)
        tbl = doc.add_table(rows=0, cols=ncol)
        tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        tbl.autofit = False
        layout = OxmlElement('w:tblLayout')
        layout.set(qn('w:type'), 'fixed')
        tbl._tbl.tblPr.append(layout)
        grid = tbl._tbl.find(qn('w:tblGrid'))
        for gc, w in zip(grid.findall(qn('w:gridCol')), widths):
            gc.set(qn('w:w'), str(int(w.twips)))
        for ri, rowdata in enumerate(block):
            row = tbl.add_row()
            for ci, celltext in enumerate(rowdata):
                cell = row.cells[ci]
                cell.width = widths[ci]
                par = cell.paragraphs[0]
                par.paragraph_format.space_after = Pt(2)
                par.paragraph_format.space_before = Pt(2)
                emit_inline(par, celltext)
                if ri == 0:
                    shade(cell, 'DCE4F0')
                    for rr in par.runs:
                        rr.bold = True
            if ri == 0:
                set_repeat_header(row)
        doc.add_paragraph()
        continue

    if stripped.startswith('- '):
        while i < len(lines) and lines[i].strip().startswith('- '):
            par = doc.add_paragraph(style='List Bullet')
            par.paragraph_format.space_after = Pt(2)
            emit_inline(par, lines[i].strip()[2:])
            i += 1
        continue

    par = doc.add_paragraph()
    emit_inline(par, stripped)
    i += 1

doc.save(OUT)
print('written', OUT)
